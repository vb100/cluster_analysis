import docx
import random
import os
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
from docx.shared import Inches, Pt
from docx.enum.dml import MSO_THEME_COLOR_INDEX

def create_document_from_template():
    return docx.Document(docx='utils/template.docx')

class PARAGRAPH_ALIGNMENT:
    LEFT = 0
    CENTER = 1
    RIGHT = 2
    JUSTIFY = 3

def document_add_png(document, png_file_path, width_inch = None, height_inch = None, alignment = PARAGRAPH_ALIGNMENT.LEFT):
    p = document.add_paragraph()
    p.alignment = alignment
    r = p.add_run()
    r.add_picture(
        png_file_path, 
        width=Inches(width_inch) if width_inch else None, 
        height=Inches(height_inch) if height_inch else None
    )


def human_format(num):
    #source: https://stackoverflow.com/questions/579310/formatting-long-numbers-as-strings/45846841#45846841
    try:
        num = int(num)
    except:
        print(f'Count not convert {num} to int. Returning 0.')
        return 0
   
    magnitude = 0
    while abs(num) >= 1000:
        magnitude += 1
        num /= 1000.0
    # add more suffixes if you need them
    result = '%.1f%s' % (num, ['', 'K', 'M', 'G', 'T', 'P'][magnitude])
    return result.replace('.0', '')


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def _is_link(text):
    return text[:4] == 'http'

def _is_nan(text):
    return text == 'nan'

def _paragraph_add_bold_text(p, text):
    r = p.add_run(text)
    r.bold = True

def add_bold_text(element, text):
    if isinstance(element, docx.table._Cell):
        _paragraph_add_bold_text(element.paragraphs[0], text)
    elif isinstance(element, docx.text.paragraph.Paragraph):
        _paragraph_add_bold_text(element, text)
    else:
        raise Exception(f'add_bold_text Unknown element type: {type(element)}')

def document_add_table(document, dataframe, format_link_label = lambda x: 'URL', autofit=True):
    rows, cols = dataframe.shape
    table = document.add_table(rows=1, cols=cols, style='Table Grid')
    header_row = table.rows[0].cells
    for i, col in enumerate(dataframe.columns):
        add_bold_text(header_row[i], col)
        header_row[i].paragraphs[0].alignment = 1 # 0 for left, 1 for center, 2 right, 3 justify ....
        
    for i, sample in dataframe.iterrows():
        row = table.add_row().cells
        for j, col in enumerate(dataframe.columns):
            if _is_link(sample[col]):
                add_hyperlink(row[j].paragraphs[0], format_link_label(sample[col]), sample[col])
            elif _is_nan(sample[col]):
                row[j].text = '-'
            else:
                row[j].text = sample[col]
         
    # Autofit columns (https://github.com/python-openxml/python-docx/issues/209)
    if autofit:
        for column in table.columns:
            for cell in column.cells:
                cell._tc.tcPr.tcW.type = 'auto'


def document_add_emotion_examples_table(document, df, top_n=20, text_col='translation', emotion_list=[]):
    assert 'emotion' in df.columns, 'This function expects a column named "emotion" in data frame.'
    assert 'emotion_score' in df.columns, 'This function expects a column named "emotion_score" in data frame.'
    
    for emotion in emotion_list:
        tmpdf = df[df['emotion'] == emotion].sort_values(by='emotion_score', ascending=False)
        samples = tmpdf[text_col]
        samples = samples[:top_n]
        document_add_table(document, pd.DataFrame({emotion.capitalize(): samples}), autofit=False)



############################################ PLOTS ##########################################
        
def create_emotion_distribution(df, top_n=5, title='', save_dir=None):
    assert 'emotion' in df.columns, 'This function expects a column named "emotion" in data frame.'
    
    my_cmap = plt.get_cmap("viridis")
    colors = [i for i in range(top_n)]
    random.shuffle(colors)
    rescale = lambda y: (y - np.min(y)) / (np.max(y) - np.min(y))
    
    emotion_vc = df['emotion'].value_counts()
    emotion_vc = emotion_vc[:top_n]

    plt.figure()
    plt.rcParams.update({'font.size': 22})
    ax = emotion_vc.plot.barh(color=my_cmap(rescale(colors)), width=0.8, figsize=(12,8), zorder=3)
    ax.invert_yaxis()
    ax.set_xlabel('Number of comments')
    ax.set_ylabel('Predicted emotion')
    plt.gca().xaxis.grid(color='gray', linestyle='dashed', zorder=0)
    if save_dir:
        plt.savefig(os.path.join(save_dir, f'{title}_emotion_barh.png'), bbox_inches="tight")
    else:
        plt.show()
    return emotion_vc.index.tolist(), emotion_vc.tolist()

import matplotlib.pyplot as plt
from datetime import datetime
from collections import Counter
import matplotlib.dates as mdates
from matplotlib.ticker import MaxNLocator
def generate_daily_trend(list_of_dates, date_format="%Y-%m-%d", xaxis_date_format='%b %d, %Y',
                   visible_only_nth=None, title='', xlabel='', ylabel='', tilt=False, savefile=None):
    date_objects = [datetime.strptime(date, date_format).date() for date in list_of_dates]
    day_counter = Counter(date_objects)
    sorted_counter = sorted(day_counter.items())
    days, frequencies = zip(*sorted_counter)

    fig, ax = plt.subplots()
    plt.gca().xaxis.set_major_formatter(mdates.DateFormatter(xaxis_date_format))
    plt.gca().xaxis.set_major_locator(mdates.DayLocator(interval=1))
    ax.bar(days, frequencies)
    ax.set_xlabel(xlabel)
    ax.set_ylabel(ylabel)
    ax.set_title(title)

    if visible_only_nth:
        for i, label in enumerate(ax.xaxis.get_ticklabels()):
            if (i + 1) % visible_only_nth != 0:
                label.set_visible(False)
    plt.grid(axis='y', linestyle='--', alpha=0.7)
    ax.yaxis.set_major_locator(MaxNLocator(integer=True))
    ax.tick_params(axis=u'y', which=u'both',length=0)
    # ax.tick_params(axis=u'x', which=u'major',length=0)
    ax.spines[['right', 'top', 'left']].set_visible(False)
    if tilt:
        plt.gcf().autofmt_xdate()
    plt.show()
    if savefile:
        plt.savefig(savefile)

