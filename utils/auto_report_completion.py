import os
from docx import Document
import pandas as pd
import os
import pickle
import requests
from bs4 import BeautifulSoup as bs
from tqdm import tqdm
import datetime
import math
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from difflib import SequenceMatcher
from docx.shared import Inches, Pt
import shutil
import random
import numpy as np
from utils.openai_utils import GPTClientForTranslation, GPTClient
from utils.doc_format import human_format, add_hyperlink, add_bold_text, document_add_table, create_document_from_template, generate_daily_trend


class FrontEndTranslator:
    def __init__(self):
        self.gpt = GPTClientForTranslation()
        self.tmp_file = '../data/prc/front_end_translations.pkl'
        if os.path.isfile(self.tmp_file):
            with open(self.tmp_file, 'rb') as fin:
                self.translations = pickle.load(fin)
        else:
            self.translations = {}

    def __call__(self, input_text):
        if type(input_text) is list:
            text_list = input_text
            text_list = [t.replace('<br>', '') for t in text_list]
            new_text_list = [t for t in text_list if t not in self.translations]
            if len(new_text_list) > 0:
                self.translations.update(self.gpt.batch_translate(new_text_list))
                with open(self.tmp_file, 'wb') as fout:
                    pickle.dump(self.translations, fout)
            return [self.translations[t] for t in text_list]
        
        elif type(input_text) is str:
            text = input_text.replace('<br>', '')
            if text not in self.translations:
                self.translations[text] = self.gpt.translate(text)
                with open(self.tmp_file, 'wb') as fout:
                    pickle.dump(self.translations, fout)
            return self.translations[text]
        
        else:
            raise Exception(f'not implemented for type {type(input_text)}')
            
front_end_translator = FrontEndTranslator()


ONE_SHOT_SELECTION = """Answer the given question by selecting N most appropriate examples. Answer using the ids of the example, in a comma separated list.

N = 3
Question: Which of the following news titles are related with Romania?
Example 0: Business Book of the Year Award 2024: winners pick their favorites
Example 1: Romania and South Korea signed the first agreement between the two countries for defense cooperation
Example 2: Rain in Bucharest on Saturday and Sunday
Example 3: Tesla cuts prices in US, China and Germany, as competition heats up
Example 4: Dacia Sandero, leader of the ranking of registrations in Europe in the first three months of 2024
Answer: 1, 2, 4

N = {}
Question: {}
{}
Answer:"""

def get_one_shot_selection_prompt(question, example_list, N=3):
    example_list = [ex.replace('\n', ' ') for ex in example_list]
    example_list = [ex.replace('"', '') for ex in example_list]
    examples = '\n'.join([f'Example {i}: {ex}' for i, ex in enumerate(example_list)])
    return ONE_SHOT_SELECTION.format(N, question, examples)


def extract_examples_for_emotion_analysis(emotion, examples):
    prompt = get_one_shot_selection_prompt(
        question = f'Which of the following YouTube comments best expres {emotion}? Please choose comments that contain more information.',
        example_list = examples,
        N = 5
        )
    gpt = GPTClient()
    rsp = gpt.get_response(prompt, model='gpt-4-turbo')
    indexes = [int(i) for i in rsp.split(', ')]
    selected_examples = [examples[i] for i in indexes]
    return selected_examples


def extract_examples_for_stance_prediction(stance, claim, examples):
    prompt = get_one_shot_selection_prompt(
        question = f'{claim} Which of the following YouTube comments best {stance.lower()} with the video? Please choose comments that contain more information.',
        example_list = examples,
        N = 5
        )
    # print(prompt)
    gpt = GPTClient()
    rsp = gpt.get_response(prompt, model='gpt-4-turbo')
    indexes = [int(i) for i in rsp.split(', ')]
    selected_examples = [examples[i] for i in indexes]
    return selected_examples


def format_link_label_youtube_video(link):
    return link.split('watch?v=')[1]
        
def document_add_top_videos_table(document, video_df, core_video_df):
    vid_dict = {'Video Id' : [], 'Theme' : [], 'Title' : []}
    for video_id, theme in zip(core_video_df.videoId.tolist(), core_video_df.theme.tolist()):
        vid_dict['Video Id'].append(f'https://www.youtube.com/watch?v={video_id}')
        vid_dict['Theme'].append(theme)
        vid_dict['Title'].append(front_end_translator(video_df[video_df.videoId == video_id].title.iloc[0]))
        
    document_add_table(document, pd.DataFrame(vid_dict), format_link_label_youtube_video)

def add_top10_video_links(document, df):
    top10videos = df.sort_values(by='commentCount', ascending=False)[:10]
    top10videos['Video ID'] = top10videos.videoId.apply(lambda x: f'https://www.youtube.com/watch?v={x}')
    top10videos['Video Title'] = front_end_translator(top10videos.title.tolist())
    document.add_heading('Interactions network', 2)
    document_add_table(document, top10videos[['Video ID', 'Video Title']], format_link_label_youtube_video)


def add_author_details(doc, users_df, video_df, top_n = 10):
    authors_df = users_df[users_df['id'].isin(video_df['channelId'].tolist())]
    
    for i in video_df['channelId'].tolist():
        if i not in authors_df['id'].tolist():
            raise Exception(f'Missing author for channel ID {i}')

    authors_df = authors_df.sort_values(by='subscriberCount', ascending=False)
    authors_df = authors_df[:top_n]
      
    top_videos = []
    for i, a in authors_df.iterrows():
        tmp_df = video_df[video_df['channelId'] == a['id']]
        top_videos.append(tmp_df['videoId'].iloc[0])
        
    
    authors_df['top_video'] = top_videos
        
    authors_df['URL'] = authors_df.id.apply(lambda x: f'https://www.youtube.com/channel/{x}')
    authors_df['Video'] = authors_df.top_video.apply(lambda x: f'https://www.youtube.com/watch?v={x}')
    authors_df['Title'] = authors_df.title
    authors_df['Country'] = authors_df.country.apply(lambda x: str(x))
    authors_df['Uploads'] = authors_df.videoCount.apply(lambda x: str(int(x)))
    authors_df['Subscribers'] = authors_df.subscriberCount.apply(lambda x: human_format(x))
    authors_df = authors_df[['URL', 'Video', 'Title', 'Country', 'Uploads', 'Subscribers']]
    
    doc.add_heading('Top video authors', 2)
    document_add_table(doc, authors_df, lambda x: 'Link', autofit=False)
    doc.add_paragraph('')    
    
def add_top_commenters(doc, comments_df, df_commenters, top_n=10):
    vc = comments_df['authorChannelId'].value_counts()
    comment_count = {user_id: comments for user_id, comments in zip(vc.index.tolist(), vc.tolist())}
    df_commenters['commentCount'] = df_commenters.id.apply(lambda x: comment_count[x] if x in comment_count else 0)
    df_commenters = df_commenters.sort_values(by='commentCount', ascending=False)
    df_commenters = df_commenters[:top_n]
    
    df_commenters['URL'] = df_commenters.id.apply(lambda x: f'https://www.youtube.com/channel/{x}')
    df_commenters['Name'] = df_commenters.title.apply(lambda x: str(x))
    df_commenters['Country'] = df_commenters.country.apply(lambda x: str(x))
    df_commenters['Uploads'] = df_commenters.videoCount.apply(lambda x: str(int(x)))
    df_commenters['Subscribers'] = df_commenters.subscriberCount.apply(lambda x: human_format(x))
    df_commenters['Comments'] = df_commenters.commentCount.apply(lambda x: str(x))
    df_commenters = df_commenters[['URL', 'Name', 'Country', 'Uploads', 'Subscribers', 'Comments']]
    
    doc.add_heading('Top video commenters', 2)
    document_add_table(doc, df_commenters, lambda x: 'Link', autofit=False)
    doc.add_paragraph('') 

    
def add_largest_discutions(doc, comm_df):
    lgdf = comm_df.sort_values(by='replyCount', ascending=False)[:10]
    lgdf['Comment'] = front_end_translator(lgdf['text'].tolist())
    lgdf['Likes'] = lgdf.likeCount.apply(lambda x: str(x))
    lgdf['Replies'] = lgdf.replyCount.apply(lambda x: str(int(x)))
    lgdf['Video URL'] = lgdf.video_id.apply(lambda x: f'https://www.youtube.com/watch?v={x}')
    lgdf['Author'] = lgdf.authorName
    lgdf['Author URL'] = f'https://www.youtube.com/channel/{lgdf.authorChannelId}'
    lgdf = lgdf[['Comment', 'Likes', 'Replies', 'Video URL', 'Author', 'Author URL']]
    
    doc.add_heading('Largest discussions', 2)
    document_add_table(doc, lgdf, lambda x: 'Link', autofit=True)
    doc.add_paragraph('') 


#################################################### TREND GENERATION ###############################################

import datetime
import calendar
import matplotlib.pyplot as plt
from matplotlib.ticker import MaxNLocator

def _add_months(sourcedate, months):
    month = sourcedate.month - 1 + months
    year = sourcedate.year + month // 12
    month = month % 12 + 1
    day = min(sourcedate.day, calendar.monthrange(year,month)[1])
    return datetime.datetime(year, month, day)

def _time_range_generator(start, end):
    while(start <= end):
        yield start
        start = _add_months(start, 1)
        
def generate_monthly_trend(df, date_col, date_format, savefile=None):
    dt_obj = pd.to_datetime(df[date_col], format=date_format)
    month = pd.Series([f'0{m}' if len(m) < 2 else m for m in dt_obj.dt.month.astype(str)])
    df['mm_yy'] = month + '-' + dt_obj.dt.year.astype(str)
    vc = df['mm_yy'].value_counts()
    data = [[month, counts] for month, counts in zip(vc, vc.index)]
    data.sort(key=lambda x: x[1])
    data = {d[1] : d[0] for d in data}
    dates = [datetime.datetime.strptime(k, "%m-%Y") for k in data]
    dates.sort()
    dates  = [d for d in _time_range_generator(dates[0], dates[-1])]
    x = [d.strftime("%m-%Y") for d in dates]
    y = [data[d] if d in data else 0 for d in x]
    x = [d.strftime("%b %Y") for d in dates]
    
    # fig = plt.figure(figsize = (18, 8))
    fig, ax = plt.subplots(figsize = (18, 8))
 
    # creating the bar plot
    ax.bar(x, y, width = 0.4)

    plt.grid(axis='y', linestyle='--', alpha=0.7)
    ax.tick_params(axis=u'y', which=u'both',length=0)
    ax.yaxis.set_major_locator(MaxNLocator(integer=True))
    ax.spines[['right', 'top', 'left']].set_visible(False)
    plt.xticks(rotation=90)
    # plt.xlabel("Month of Year")
    # plt.ylabel("No. of videos")
 
    if savefile:
        plt.savefig(savefile, bbox_inches='tight')
    
    
def generate_and_add_trend(document, df, date_format='%d-%m-%y %M:%S', date_col='publishedAtSQL', binning='monthly'):
    assert binning in ['monthly', 'daily']
    if binning == 'monthly':
        generate_monthly_trend(df, date_col=date_col, date_format=date_format, savefile='../data/prc/trend.png')
    elif binning == 'daily':
        list_of_dates = df[date_col].tolist()
        generate_daily_trend(list_of_dates, date_format=date_format, xaxis_date_format='%b %d, %Y',
                   visible_only_nth=None, title='', xlabel='', ylabel='', tilt=True, savefile='../data/prc/trend.png')
        
    document.add_heading('Publish Trend', 2)
    p = document.add_paragraph()
    p.alignment = 1
    r = p.add_run()
    r.add_picture('../data/prc/trend.png', width=Inches(6.4), height=Inches(3.2))


########################################################## SUMMARIZATION ##########################################################
SUMMARY_TEMPLATE = """Summarize the transcript of a YouTube video in a concise format. 
The summary should begin with the main idea of the content, encapsulated in 30 to 50 words. 
The summary should continue with five bullet points highlighting crucial information from the transcript.
### Transcript:
"{}"

### Summary:
"""
SYSTEM_PROMPT = "You are a summarization bot."

from utils.openai_utils import GPTClient
gpt = GPTClient()


def generate_and_save_summary(video_id, transcript_df, save_dir='../data/prc/summaries/'):
    outfile_path = os.path.join(save_dir, f"{video_id}_summary.txt")

    if os.path.isfile(outfile_path):
        print(f'Summary for transcript {video_id} already exists. Skipping this step...')
        return


    text = transcript_df[transcript_df.video_id == video_id].content.iloc[0]
    text = text[:8000]
    text = text.replace('"', '')
    prompt = SUMMARY_TEMPLATE.format(text)
    response = gpt.get_response(prompt, system_prompt=SYSTEM_PROMPT, temperature=0.5, max_tokens=512)
    with open(outfile_path, 'wt') as fout:
        fout.write(response)


def generate_top5_transcript_summary(df):
    top5videos = df.sort_values(by='commentCount', ascending=False)[:5]
    for i, row in tqdm(top5videos.iterrows(), total=len(top5videos), desc='Generting summaries'):
        generate_and_save_summary(row["videoId"])


######################################################### THUMBNAILS ###############################################
def download_thumbnail(video_id, save_dir='../data/prc/thumbnails/'):
    outfile_path = os.path.join(save_dir, f'{video_id}_thumbnail.png')
    if os.path.isfile(outfile_path):
        print(f'Thumbnail for video {video_id} already exists. Skipping this step...')
        return

    response = requests.get(f'https://img.youtube.com/vi/{video_id}/hqdefault.jpg', stream=True)
    with open(outfile_path, 'wb') as out_file:
        shutil.copyfileobj(response.raw, out_file)
    
    del response

def download_top5_video_thumbnail(df):
    top5videos = df.sort_values(by='commentCount', ascending=False)[:5]
    for i, row in top5videos.iterrows():
        download_thumbnail(row["videoId"])


from wordcloud import WordCloud, STOPWORDS
import matplotlib.pyplot as plt
import pandas as pd
import argparse

def get_word_cloud(word_list, save_file_name):
    w, h = 16, 8
    comment_words = ''
    stopwords = set(STOPWORDS)
    
    # iterate through the csv file
    for val in word_list:

        # typecaste each val to string
        val = str(val)

        # split the value
        tokens = val.split()

        # Converts each token into lowercase
        for i in range(len(tokens)):
            tokens[i] = tokens[i].lower()

        comment_words += " ".join(tokens)+" "

    wordcloud = WordCloud(width = w * 100, height = h * 100,
                    background_color ='white',
                    collocations = False,
                    stopwords = stopwords,
                    min_font_size = 10).generate(comment_words)

    # plot the WordCloud image                      
    plt.figure(figsize = (w, h), facecolor = None)
    plt.imshow(wordcloud)
    plt.axis("off")
    plt.tight_layout(pad = 0)
    plt.savefig(save_file_name)
    
def flatten(l):
    return [item for sublist in l for item in sublist]

def process_word_list(word_list):
    word_list = [w.split(';') for w in word_list]
    word_list = flatten(word_list)
    word_list = [w.strip() for w in word_list]
    return word_list 

def generate_word_clouds(comments_df):
    targets = comments_df['targets'].dropna().tolist()
    opinions = comments_df['opinions'].dropna().tolist()

    targets = process_word_list(targets)
    opinions = process_word_list(opinions)

    get_word_cloud(targets, f'../report/targets_word_cloud.png')
    get_word_cloud(opinions, f'../report/opinions_word_cloud.png')

def add_targets_and_opinions_word_clouds(document):
    document.add_heading('Opinion mining', 2)
    p = document.add_paragraph()
    p.alignment = 1
    r = p.add_run()
    r.add_picture('../report/targets_word_cloud.png', width=Inches(6.4), height=Inches(3.2))
    r.add_text('Targets from all comments')
    
    document.add_paragraph(' ')
    
    p = document.add_paragraph()
    p.alignment = 1
    r = p.add_run()
    r.add_picture('../report/opinions_word_cloud.png', width=Inches(6.4), height=Inches(3.2))
    r.add_text('Opinions from all comments')   


def add_opinion_mining_examples(document, comments_df, targets):
    table = document.add_table(rows=1, cols=3, style='Table Grid')
    
    # Adding heading in the 1st row of the table
    row = table.rows[0].cells
    row[0].text = 'Opinion target'
    row[1].text = 'Opinion'
    row[2].text = 'Actual comment'
    
    for t in targets:
        tmp_df = comments_df[comments_df['targets'].str.contains(t, na=False)]

        for i, sample in tmp_df[:10].iterrows():
            # Adding a row and then adding data in it.
            row = table.add_row().cells
            # Converting id to string as table can only take string input
            row[0].text = str(sample['targets'])
            row[1].text = str(sample['opinions'])
            row[2].text = str(sample['text'])
            # row[2].text = str(front_end_translator(sample['text']))
    document.add_paragraph('')


# ENTITY TABLE

def get_matching_seq(string1, string2):
    match = SequenceMatcher(None, string1, string2).find_longest_match(0, len(string1), 0, len(string2))
    return string2[match.b:match.b + match.size]


def remove_closely_related_entities(seq):
    changes = {}
    print('Removing closely related entities...')
    while True:
        pbar = tqdm(enumerate(seq), total=len(seq))
        for i, s1 in pbar:
            for j in range(i+1, len(seq)):
                s2 = seq[j]

                if s1 == s2:
                    continue

                matching_seq = get_matching_seq(s1.lower(), s2.lower())
                if len(s1) > len(s2):
                    if len(matching_seq) / len(s1) >= 0.75:
                        changes[i] = s2
                        
                if len(s2) > len(s1):
                    if len(matching_seq) / len(s2) >= 0.75:
                        changes[j] = s1
                        
                if len(s1) == len(s2):
                    if len(matching_seq) / len(s1) >= 0.75 or len(matching_seq) / len(s2) >= 0.75:
                        if s1 < s2:
                            changes[i] = s2
                        else:
                            changes[j] = s1
                        pbar.set_description(f'"{s1}" vs "{s2}"')

        if len(changes) == 0:
            break
                        
        for i, change in changes.items():
            ## 'Replacing {seq[i]} with {change}'
            seq[i] = change
        
        changes = {}
    
    return seq

def add_entity_table(doc, comments):
    table_dict = {}
    for col in ['ent_Organizations', 'ent_Products', 'ent_Persons', 'ent_LAW']:
        tmp = comments[col][comments[col].notnull()]
        lst = tmp.value_counts()[:500].index.tolist()
        counts = tmp.value_counts()[:500].tolist()
        lst = remove_closely_related_entities(lst)
        big_lst = [el for ent, count in zip(lst, counts) for el in [ent] * count]
        # comments.loc[comments[col].notnull(), col] = remove_closely_related_entities(tmp)
        # tmp = comments[col][comments[col].notnull()]
        vc = pd.Series(big_lst).value_counts()
        vc = [(i, c)for i, c in zip(vc.index, vc)]
        vc.sort(key=lambda x: x[1], reverse=True)
        table_dict[col] = vc[:20]

    # Removing duplicates between columns    
    for col1 in ['ent_Products', 'ent_Persons', 'ent_Organizations', 'ent_LAW']:
        for col2 in table_dict:
            if col1 == col2:
                continue

            to_remove1 = []
            to_remove2 = []
            for i, ent1 in enumerate(table_dict[col1]):
                for j, ent2 in enumerate(table_dict[col2]):
                    if ent1[0] == ent2[0]:
                        if ent1[1] >= ent2[1]:
                            print(f'Removing {ent2} from {col2}, keeping {ent1} from {col1}')
                            to_remove2.append(j)
                        else:
                            print(f'Removing {ent1} from {col1}, keeping {ent2} from {col2}')
                            to_remove1.append(i)
            for i in sorted(to_remove1, reverse=True):
                del table_dict[col1][i]

            for j in sorted(to_remove2, reverse=True):
                del table_dict[col2][j]    

    # Creating a table object
    doc.add_heading('Entity extraction', 2)
    
    ent_table_dict = {}
    for k in ['ent_Organizations', 'ent_Products', 'ent_Persons', 'ent_LAW']:
        k2 = k[4:].lower().capitalize()
        ent_table_dict[k2] = [e[0] for e in table_dict[k]][:15]
        ent_table_dict[k2] += [''] * (15 - len(ent_table_dict[k2]))

    document_add_table(doc, pd.DataFrame(ent_table_dict))
    doc.add_paragraph('')  


def add_entity_extraction_examples(document, comments_df, targets):
    for t in targets:
        tmp_df = comments_df[comments_df['translation'].str.contains(t, na=False, case=False)]
        doc_df = pd.DataFrame({t : front_end_translator(tmp_df['text'][:10].tolist())})
        document_add_table(document, doc_df)
        document.add_paragraph('')

# Keyword detection
def add_keyword_detection_result(document, input_dir):
    document.add_heading('YouTube video keyword detection', 2)    

    p = document.add_paragraph('We further show messages with similar meaning which can be found in multiple videos.')

    keyword_cluster_files = [file for file in  os.listdir(input_dir) if 'keyword_cluster_' in file]
    keyword_cluster_files.sort(key=lambda x: int(x[16:-4]))

    for file in keyword_cluster_files:
        tmpdf = pd.read_csv(os.path.join(input_dir, file))


        tmpdict = {}
        for i, row in tmpdf.iterrows():
            if row['videos'] not in tmpdict:
                tmpdict[row['videos']] = [row['messages'], 1]
            else:
                if len(tmpdict[row['videos']][0]) < len(row['messages']):
                    tmpdict[row['videos']][0] = row['messages']
                tmpdict[row['videos']][1] += 1

        tmplist = [[video, msg, n] for video, (msg, n) in tmpdict.items()]
        tmplist.sort(key=lambda x: x[2], reverse=True)

            # Table data in a form of list
        table = document.add_table(rows=2, cols=3, style='Table Grid')
        # Adding heading in the 1st row of the table
        row = table.rows[0].cells
        row[0].text = f'Important message: x ({len(tmplist)} videos)'

        row = table.rows[1].cells
        row[0].text = 'Video Message'
        row[1].text = 'Mentions'
        row[2].text = 'Video ID'

        for video_id, msg, n in tmplist:

            # Adding a row and then adding data in it.
            row = table.add_row().cells
            # Converting id to string as table can only take string input
            row[0].text = str(msg)
            row[1].text = str(n)
            add_hyperlink(row[2].paragraphs[0], f'{video_id}', f'https://www.youtube.com/watch?v={video_id}')

        document.add_paragraph('')

# Video summary
def human_format_date(date_str) : 
    for fmt in ('%Y-%m-%d %M:%S', '%Y-%m-%d %H:%M:%S', '%d/%m/%Y %M:%S', "%Y-%m-%dT%H:%M:%SZ"):
        try:
            return datetime.datetime.strptime(date_str, fmt).strftime('%B %d, %Y')
        except ValueError:
            pass
    raise ValueError(f'no valid date format found. Input value: {date_str}')

def document_add_video_summary_and_thumbnail(document, video_id, video_df, users_df, summary_save_dir='../data/prc/summaries/', thumbnail_save_dir='../data/prc/thumbnails/'):
    video_row = video_df[video_df['videoId'] == video_id].iloc[0]
    channel_id = video_row['channelId']
    channel_row = users_df[users_df['id'] == channel_id].iloc[0]

    # adding info paragraph
    p = document.add_paragraph(f'Video ')
    add_hyperlink(p, f'{video_id}', f'https://www.youtube.com/watch?v={video_id}')
    p.add_run(f': ')

    r = p.add_run(f'"{video_row["title"]}" ({front_end_translator(video_row["title"])}),')
    r.italic = True

    p.add_run(f'{human_format(video_row["viewCount"])} views, '
                f'{human_format(video_row["likeCount"])} likes, '
                f'{human_format(video_row["commentCount"])} comments. ')

    p.add_run(f'Uploaded on {human_format_date(video_row["published"])}. '
                f'Author: {channel_row["title"]}, {human_format(channel_row["subscriberCount"])} subscribers, '
                f'{human_format(channel_row["videoCount"])} uploads.')
    
    p = document.add_paragraph()
    p.alignment = 1
    r = p.add_run()
    r.add_picture(os.path.join(thumbnail_save_dir, f'{video_id}_thumbnail.png'), width=Inches(4), height=Inches(3))
    
    with open(os.path.join(summary_save_dir, f'{video_id}_summary.txt'), 'rt') as fin:
        summary = fin.read()
        summary_lines = summary.split('\n')
        p = document.add_paragraph()
        r = p.add_run('Content summary: ')
        r.bold = True
        p.add_run(summary_lines[0])
        for line in summary_lines[1:]:
            document.add_paragraph(line) 


# Emotion analysis
def document_add_video_emotion_analysis(doc, video_id, comments_df, emotion_list=['curiosity', 'confusion'], add_top5_to_emotion_list=False, use_gpt_selection=False):
    doc.add_heading('Emotion Analysis', 2)
    my_cmap = plt.get_cmap("viridis")
    rescale = lambda y: (y - np.min(y)) / (np.max(y) - np.min(y))
    
    # Generating emotion horizontal bar graph
    tmpdf = comments_df[comments_df['video_id'] == video_id]
    vc = tmpdf['emotion'].value_counts()
    vc = vc[:5]
    
    if add_top5_to_emotion_list:
        emotion_list += vc.index.tolist()
        emotion_list = list(set(emotion_list))

    colors = [i for i in range(5)]
    random.shuffle(colors)
    plt.figure()
    plt.rcParams.update({'font.size': 22})
    ax = vc.plot.barh(color=my_cmap(rescale(colors)), width=0.8, figsize=(12,8), zorder=3)
    ax.invert_yaxis()
    ax.set_xlabel('Number of comments')
    ax.set_ylabel('Predicted emotion')
    plt.gca().xaxis.grid(color='gray', linestyle='dashed', zorder=0)
    plt.savefig(f'../data/prc/{video_id}_emotion_barh.png', bbox_inches="tight")
    plt.show()

    # Adding thumbnail + bar graph table
    table = doc.add_table(rows=1, cols=2, style='Table Grid')
    row = table.rows[0].cells
    r = row[0].paragraphs[0].add_run()
    r.add_picture(f'../data/prc/thumbnails/{video_id}_thumbnail.png', width=Inches(2.6), height=Inches(2))
    r = row[1].paragraphs[0].add_run()
    r.add_picture(f'../data/prc/{video_id}_emotion_barh.png', width=Inches(3), height=Inches(2))
    

    for emotion in emotion_list:
        samples = tmpdf[tmpdf['emotion'] == emotion].sort_values(by='emotion_score', ascending=False)
        samples = samples['text'].tolist()[:20]
        samples = front_end_translator(samples)
        if use_gpt_selection:
            samples = extract_examples_for_emotion_analysis(emotion, samples)
        document_add_table(doc, pd.DataFrame({emotion.capitalize() : samples}))


def document_add_video_main_discution_themes_antet(document, video_id, video_df):
    video_row = video_df[video_df['videoId'] == video_id]
    document.add_heading(f'Main discussion themes', 2)


    p = document.add_paragraph(f'Video ')
    add_hyperlink(p, f'{video_id}', f'https://www.youtube.com/watch?v={video_id}')
    p.add_run(f'. Title: {video_row["title"].iloc[0]} ({front_end_translator(video_row["title"].iloc[0])}). Comments: {int(video_row["commentCount"].iloc[0])}.')
    
    p = document.add_paragraph('')
    r = p.add_run()
    r.add_picture(f'../data/prc/thumbnails/{video_id}_thumbnail.png', width=Inches(1.2), height=Inches(0.9))
    p = document.add_paragraph('')


def document_add_comments_triggered(document, video_id, video_triggering_comms):
    document.add_heading('YouTube Content Analysis', 2)    

            
    triggering_comms = video_triggering_comms[video_id]
    # Table data in a form of list
    data = [(t[0], t[1]) for t in triggering_comms[:10]]

    # Creating a table object
    table = document.add_table(rows=1, cols=3, style='Table Grid')

    # Adding heading in the 1st row of the table
    row = table.rows[0].cells
    row[0].text = 'Message'
    row[1].text = 'Nr Comments'
    row[2].text = 'Video '
    add_hyperlink(row[2].paragraphs[0], f'{video_id}', f'https://www.youtube.com/watch?v={video_id}')
    r = row[2].paragraphs[0].add_run()
    r.add_picture(f'../data/prc/thumbnails/{video_id}_thumbnail.png', width=Inches(2), height=Inches(1.5))
    
    video_text = f'Video https://www.youtube.com/watch?v={video_id}'
    # Adding data from the list to the table
    for i, (msg, nr_comments) in enumerate(data):

        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(msg)
        row[1].text = str(nr_comments)

    document.add_paragraph('')


# Stance prediction
def show_pie(tmp_df):
    colors = {
        'AGREE' : 'green',
        'DISAGREE' : 'red',
        'UNRELATED' : 'grey'
    }
    tmp_df = tmp_df[tmp_df['stance'] != 'Inappropriate prompt!']
    vc = tmp_df['stance'].value_counts()
    vc.plot.pie(explode=[0.03] * len(vc), colors=[colors[l] for l in vc.index.tolist()], label='')


def document_add_stance_prediction(document, video_id, stance_prediction_df, use_gpt_selection=False, select_top_k = 5, save_dir='../data/prc/stance_prediction/'):
    # claims_df = pd.read_csv('../data/prc/stance_prediction.csv')

    document.add_heading('Stance Prediction', 2)

    p = document.add_paragraph(f'Video ')
    add_hyperlink(p, f'{video_id}', f'https://www.youtube.com/watch?v={video_id}')

    p = document.add_paragraph('')
    r = p.add_run()
    r.add_picture(f'../data/prc/thumbnails/{video_id}_thumbnail.png', width=Inches(1.2), height=Inches(0.9))

    claims = stance_prediction_df['claim'].unique().tolist()
    for j, claim in enumerate(claims):
        tmp_claim_df = stance_prediction_df[stance_prediction_df['claim'] == claim]
        p = document.add_paragraph('')
        add_bold_text(p, f'Claim {j + 1}: {claim}')
        
        # claim pie chart
        plt.figure()
        show_pie(tmp_claim_df)
        plt.savefig(os.path.join(save_dir, f'{video_id}_claim_{j}_pie.png'), bbox_inches="tight")
        plt.show()
        p = document.add_paragraph('')
        r = p.add_run()
        r.add_picture(os.path.join(save_dir, f'{video_id}_claim_{j}_pie.png'), height=Inches(1.5))
        document.add_paragraph('')
        
        # claim table
        agree_list = tmp_claim_df[tmp_claim_df['stance'] == 'AGREE']['comment'].tolist()
        disagree_list = tmp_claim_df[tmp_claim_df['stance'] == 'DISAGREE']['comment'].tolist()
        
        
        agree_list.sort(key=lambda x: len(x), reverse=True)
        disagree_list.sort(key=lambda x: len(x), reverse=True)
        
        if len(agree_list) > select_top_k and use_gpt_selection:
            agree_list = extract_examples_for_stance_prediction('AGREE', claim, agree_list)

        if len(disagree_list) > select_top_k and use_gpt_selection:
            disagree_list = extract_examples_for_stance_prediction('DISAGREE', claim, disagree_list)

        max_len = max([len(agree_list), len(disagree_list)])
        agree_list += [''] * (max_len - len(agree_list))
        disagree_list += [''] * (max_len - len(disagree_list))
        
        agreement_df = pd.DataFrame({'AGREE' : agree_list[:select_top_k], 'DISAGREE' : disagree_list[:select_top_k]})
        document_add_table(document, agreement_df)


# Transcripts

def document_add_video_transcripts(document, video_id, video_df, transcript_df, thumbnail_dir='../data/prc/thumbnails/'):
    video_row = video_df[video_df['videoId'] == video_id]    
    p = document.add_paragraph('')
    r = p.add_run()
    r.add_picture(os.path.join(thumbnail_dir, f'{video_id}_thumbnail.png'), width=Inches(1.2), height=Inches(0.9))
    
    document.add_paragraph(front_end_translator(video_row['title'].iloc[0]))
    

    transcript = transcript_df[transcript_df['video_id'] == video_id]['content'].iloc[0]
    if len(transcript) > 10900:
        transcript = transcript[:10900] + ' [...]'
    
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    row = table.rows[0].cells
    row[0].text = transcript.strip()
    font = row[0].paragraphs[0].runs[0].font
    font.size = Pt(7)
    
    document.add_paragraph('')


def flatten(l):
    return [item for sublist in l for item in sublist]

def show_report_stats(comments_df):
    print('=' * 80)
    print('Top emotions in the report:')
    print(comments_df['emotion'].value_counts()[:10])
    print('=' * 80)

    def process_word_list(word_list):
        word_list = [w.split(';') for w in word_list]
        word_list = flatten(word_list)
        word_list = [w.strip() for w in word_list]
        return word_list

    target_list = comments_df['targets'].dropna().tolist()
    target_list = process_word_list(target_list)
    print('Top targets in Opinion Mining:')
    print(pd.Series(target_list).value_counts()[:10])
    print('=' * 80)

    print('Top entities in the report:')
    entities = {}
    for col in ['ent_Products', 'ent_Persons', 'ent_Organizations', 'ent_LAW']:
        ent_list = comments_df[col].dropna().tolist()
        vc = pd.Series(ent_list).value_counts()[:10]
        print( '-' * 20 +  f"{col}" + '-' * 20)
        print(vc)

